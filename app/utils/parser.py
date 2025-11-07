import io
import pdfplumber
import tempfile
import os
from docx2pdf import convert
from docx import Document
from collections import Counter
from app.utils.parserUtils import ParserUtils
from app.models.resume import ( 
    ResumeDocument,
    ResumeEntry,
    ResumeSection,
    ResumeMetadata)


class Parser:

    @staticmethod
    def pdfParser(file_bytes: bytes) -> ResumeDocument:
        """
        Parse PDF and return a ResumeDocument with sections + metadata.
        Each word is turned into a ResumeEntry with formatting metadata.
        """
        entries: list[ResumeEntry] = []
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page_idx, page in enumerate(pdf.pages):
                    words = page.extract_words(use_text_flow=True)
                    lines_map = {}

                    for w in words:
                        line_y = round(w["top"])
                        lines_map.setdefault(line_y, []).append(w)

                    for _, (y, word_list) in enumerate(sorted(lines_map.items(), key=lambda x: x[0])):
                        line_words = []
                        line_top = float("inf")
                        line_bottom = float("-inf")

                        for w in word_list:
                            line_words.append(w["text"])
                            line_top = min(w["top"], line_top)
                            line_bottom = max(w["bottom"], line_bottom)


                        line_text = " ".join(line_words).strip()
                        if not line_text:
                            continue

                        line_chars = [
                            c for c in page.chars
                            if line_top -1 <= c["top"] <= line_bottom + 1
                        ]

                        if line_chars:
                            fonts = [c.get("fontname", "") for c in line_chars if c.get("fontname")]
                            sizes = [c.get("size") for c in line_chars if c.get("size")]
                            colors = [c.get("non_stroking_color") for c in line_chars if c.get("non_stroking_color")]
                            fontname = Counter(fonts).most_common(1)[0][0] if fonts else None
                            size = Counter(sizes).most_common(1)[0][0] if sizes else None
                            color = Counter(colors).most_common(1)[0][0] if colors else None
                        else:
                            fontname, size, color = None, None, None

                        entry = ResumeEntry(
                            id=f"pdf_{page_idx}_{len(entries)}",
                            content=line_text,
                            font=fontname,
                            font_size=str(size) if size else None,
                            font_color=str(color) if color else None,
                            bold="Bold" in (fontname or ""),
                            italic="Italic" in (fontname or ""),
                            bullet=line_text.startswith("•", "-", "◦"),
                            bullet_symbol=line_text[0] if line_text.startswith(("•", "-", "◦")) else None,
                            technical_skills=set()
                        )
                        if entry.content.strip():
                            entries.append(entry)

                metadata = ResumeMetadata(
                    pages=len(pdf.pages),
                    columns=1,
                    border=False,
                    alignment="left",
                )

                sections = ParserUtils.build_resume_sections(entries)

                doc = ResumeDocument(
                    sections=sections,
                    metadata=metadata,
                    entries=entries  # optional flat cache
                )
                return doc

        except Exception as e:
            raise RuntimeError(f"Error parsing PDF: {e}")

    @staticmethod
    def docParser(file_bytes: bytes) -> ResumeDocument:
        """
        Parse DOCX and return a ResumeDocument with sections + metadata.
        Captures text + formatting for each run.
        """
        entries: list[ResumeEntry] = []
        try:
            doc = Document(io.BytesIO(file_bytes))
            for para_idx, para in enumerate(doc.paragraphs):
                if not para.text.strip():
                    continue
                
                line_text = para.text.strip()
                fonts, sizes, colors = [], [], []
                is_bold, is_italic = False, False

                for run in para.runs:
                    if run.text.strip():
                        if run.font and run.font.name:
                            fonts.append(run.font.name)
                        if run.font and run.font.size:
                            sizes.append(run.font.size.pt)
                        if run.font and run.font.color and run.font.color.rgb:
                            colors.append(str(run.font.color.rgb))
                        if run.bold:
                            is_bold = True
                        if run.italic:
                            is_italic = True
                
                fontname = Counter(fonts).most_common(1)[0][0] if fonts else None
                size = Counter(sizes).most_common(1)[0][0] if sizes else None
                color = Counter(colors).most_common(1)[0][0] if colors else None

                bullet_symbol = None
                bullet = False
                if para.style and "List" in para.style.name:
                    bullet = True
                    bullet_symbol = "•" if "Bullet" in para.style.name else None
                elif line_text.startswith(("•", "-", "◦")):
                    bullet = True
                    bullet_symbol = line_text[0]

                entry = ResumeEntry(
                    id=f"docx_{para_idx}_{len(entries)}",
                    content=line_text,
                    font=fontname,
                    font_size=str(size) if size else None,
                    font_color=str(color) if color else None,
                    bold=is_bold,
                    italic=is_italic,
                    bullet=bullet,
                    bullet_symbol=bullet_symbol,
                    technical_skills=set()
                )

                entries.append(entry)

            metadata = ResumeMetadata(
                pages=1,  # DOCX doesn’t expose page count directly
                columns=1,
                border=False,
                alignment="left",
            )

            sections = ParserUtils.build_resume_sections(entries)

            doc = ResumeDocument(
                sections=sections,
                metadata=metadata,
                entries=entries  # optional flat cache
            )
            return doc

        except Exception as e:
            raise RuntimeError(f"Error parsing DOCX: {e}")
        
    
    @staticmethod
    def docParser_2(file_bytes: bytes) -> ResumeDocument:
        """
        Converts DOCX to PDF internally and parses it using the same pdfParser logic.
        Returns a unified ResumeDocument.
        """
        from app.utils.parser import Parser  # import inside to avoid circular deps

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
            temp_docx.write(file_bytes)
            temp_docx_path = temp_docx.name

        temp_pdf_path = temp_docx_path.replace(".docx", ".pdf")

        try:
            # Convert DOCX → PDF
            convert(temp_docx_path, temp_pdf_path)

            # Read the generated PDF
            with open(temp_pdf_path, "rb") as f:
                pdf_bytes = f.read()

            # Parse using existing pdfParser
            resume_doc = Parser.pdfParser(pdf_bytes)
            return resume_doc

        except Exception as e:
            raise RuntimeError(f"Error converting DOCX to PDF: {e}")

        finally:
            # Clean up temp files
            for path in [temp_docx_path, temp_pdf_path]:
                if os.path.exists(path):
                    os.remove(path)
